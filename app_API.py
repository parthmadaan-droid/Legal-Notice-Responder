from fastapi import FastAPI, HTTPException, File, UploadFile, Form, Depends, BackgroundTasks, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
from typing import Optional, Dict, Any, List, Union
import os
import io
import uuid
import json
import re
import base64
from datetime import datetime, timedelta
import logging
import time
import asyncio
from functools import partial
import concurrent.futures
import threading
import shutil
import tempfile

# Document processing libraries
from dateutil import parser
from dateutil.relativedelta import relativedelta
from PIL import Image
import pytesseract

# OpenAI integration
import openai
from dotenv import load_dotenv

# PDF handling
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('contract_analyzer_api.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()
logger.info("Application started - Loading environment variables")

# Setup OpenAI
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
if not OPENAI_API_KEY:
    logger.error("OpenAI API key not found in environment variables")
else:
    openai.api_key = OPENAI_API_KEY
    logger.info("OpenAI API key loaded successfully")

# Initialize FastAPI
app = FastAPI(
    title="Multi-Document Contract Analyzer API",
    description="API for analyzing and extracting information from various contract types",
    version="1.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Adjust in production to specific domains
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Temporary file storage (in-memory for API)
# In a production environment, you would use a database or object storage
DOCUMENT_STORE = {}
ANALYSIS_RESULTS = {}
SESSIONS = {}

# Create a temporary directory for file processing
TEMP_DIR = tempfile.mkdtemp()
logger.info(f"Created temporary directory: {TEMP_DIR}")

# Pydantic models
class DocumentUploadResponse(BaseModel):
    document_id: str
    filename: str
    document_type: str
    status: str
    message: str

class AnalysisResponse(BaseModel):
    document_id: str
    document_type: str
    contract_summary: Optional[str] = None
    status: str
    extraction_method: Optional[str] = None
    dates: Optional[Dict[str, Any]] = None
    analysis_complete: bool = False
    error: Optional[str] = None

class QuestionRequest(BaseModel):
    document_id: str
    question: str
    session_id: Optional[str] = None

class QuestionResponse(BaseModel):
    document_id: str
    question: str
    answer: str
    session_id: str
    processing_time_ms: float

class DocumentInfoResponse(BaseModel):
    document_id: str
    document_type: str
    filename: str
    extraction_method: Optional[str] = None
    char_count: int = 0
    word_count: int = 0
    status: str
    analysis_complete: bool = False

# Helper Functions
def setup_openai_client():
    """Setup OpenAI client using environment variables"""
    if not OPENAI_API_KEY:
        logger.error("OpenAI API key not found in environment variables")
        raise HTTPException(
            status_code=500, 
            detail="OpenAI API key not configured. Please set OPENAI_API_KEY in .env file."
        )
    
    # Try to use the newer client interface if available
    try:
        from openai import OpenAI
        client = OpenAI(api_key=OPENAI_API_KEY)
        logger.info("Using new OpenAI client interface")
        return client
    except ImportError:
        # Fallback to older interface
        logger.info("Using legacy OpenAI interface")
        return None

def get_extraction_schema(document_type: str) -> str:
    """Get the appropriate JSON schema based on document type"""
    
    if document_type == "Rental":
        return """
        {
            "document_type": "Rental Agreement",
            "contract_summary": "Brief 2-3 sentence summary of the rental contract",
            "parties": {
                "landlord": "Name and details",
                "tenant": "Name and details"
            },
            "property_details": {
                "address": "Property address",
                "type": "Property type",
                "description": "Property description"
            },
            "financial_terms": {
                "monthly_rent": "Monthly rent amount",
                "security_deposit": "Security deposit amount",
                "other_fees": "Other fees and charges"
            },
            "dates": {
                "start_date": "YYYY-MM-DD format",
                "end_date": "YYYY-MM-DD format",
                "lease_term": "Duration in months/years"
            },
            "key_terms": ["List of important terms and conditions"],
            "special_clauses": ["Any special clauses or restrictions"]
        }"""
    
    elif document_type == "NDA":
        return """
        {
            "document_type": "Non-Disclosure Agreement",
            "contract_summary": "Brief 2-3 sentence summary of the NDA",
            "parties": {
                "disclosing_party": "Party disclosing confidential information",
                "receiving_party": "Party receiving confidential information"
            },
            "confidentiality_scope": {
                "definition": "What constitutes confidential information",
                "exclusions": "What is excluded from confidentiality",
                "purpose": "Purpose for which information can be used"
            },
            "obligations": {
                "non_disclosure": "Non-disclosure obligations",
                "non_use": "Restrictions on use of information",
                "return_destroy": "Obligations to return or destroy information"
            },
            "dates": {
                "execution_date": "YYYY-MM-DD format",
                "effective_date": "YYYY-MM-DD format", 
                "duration": "Duration of confidentiality obligations",
                "survival_period": "How long obligations survive termination"
            },
            "restrictions": ["List of specific restrictions and limitations"],
            "remedies": ["Available remedies for breach"]
        }"""
    
    elif document_type == "MSA":
        return """
        {
            "document_type": "Master Service Agreement",
            "contract_summary": "Brief 2-3 sentence summary of the MSA",
            "parties": {
                "service_provider": "Company providing services",
                "client": "Company receiving services"
            },
            "service_details": {
                "description": "Description of services to be provided",
                "scope": "Scope of services",
                "deliverables": "Key deliverables"
            },
            "financial_terms": {
                "pricing_model": "How services are priced",
                "payment_terms": "Payment schedule and terms",
                "invoicing": "Invoicing procedures"
            },
            "dates": {
                "execution_date": "YYYY-MM-DD format",
                "effective_date": "YYYY-MM-DD format",
                "initial_term": "Initial contract term",
                "renewal_terms": "Renewal provisions"
            },
            "termination": {
                "termination_rights": "How either party can terminate",
                "notice_period": "Required notice period",
                "effect_of_termination": "What happens upon termination"
            },
            "key_obligations": ["List of key obligations for each party"],
            "liability_indemnity": ["Limitation of liability and indemnification terms"]
        }"""
    
    elif document_type == "Insurance":
        return """
        {
            "document_type": "Insurance Policy",
            "contract_summary": "Brief 2-3 sentence summary of the insurance policy",
            "parties": {
                "insurer": "Insurance company name and details",
                "policyholder": "Insured person/entity name and details",
                "beneficiary": "Beneficiary details if applicable"
            },
            "policy_details": {
                "policy_number": "Policy identification number",
                "insurance_type": "Type of insurance (Health/Auto/Life/Property/etc.)",
                "coverage_amount": "Coverage limits and amounts",
                "deductible": "Deductible amount"
            },
            "financial_terms": {
                "premium_amount": "Premium cost and payment frequency",
                "copayment": "Copayment amounts if applicable",
                "coinsurance": "Coinsurance percentages if applicable",
                "out_of_pocket_maximum": "Maximum out-of-pocket costs"
            },
            "dates": {
                "policy_start": "YYYY-MM-DD format",
                "policy_end": "YYYY-MM-DD format",
                "renewal_date": "YYYY-MM-DD format",
                "grace_period": "Grace period for payments"
            },
            "coverage_details": ["List of what is covered by the policy"],
            "exclusions": ["List of what is excluded from coverage"],
            "claim_procedures": ["How to file claims and claim process"]
        }"""
    
    elif document_type == "MOU":
        return """
        {
            "document_type": "Memorandum of Understanding",
            "contract_summary": "Brief 2-3 sentence summary of the MOU",
            "parties": {
                "party_1": "First organization/entity name and details",
                "party_2": "Second organization/entity name and details",
                "additional_parties": "Other parties if multi-party MOU"
            },
            "purpose": {
                "objective": "Main objective and purpose of the MOU",
                "background": "Background context for the agreement",
                "scope": "Scope and areas covered by the MOU"
            },
            "responsibilities": {
                "party_1_obligations": "First party's duties and responsibilities",
                "party_2_obligations": "Second party's duties and responsibilities",
                "shared_responsibilities": "Joint or shared obligations"
            },
            "dates": {
                "execution_date": "YYYY-MM-DD format",
                "effective_date": "YYYY-MM-DD format",
                "duration": "Duration or term of the MOU",
                "review_date": "YYYY-MM-DD format for periodic review"
            },
            "key_terms": ["List of important terms and conditions"],
            "termination_conditions": ["Conditions under which MOU can be terminated"],
            "governance": ["How the MOU will be managed and monitored"]
        }"""
    
    else:
        return get_extraction_schema("Rental")  # Default to rental

async def extract_text_from_file(file_path: str, file_type: str, file_name: str) -> Optional[str]:
    """Extract text from various file formats"""
    logger.info(f"Starting text extraction from file: {file_name}")
    logger.info(f"File type: {file_type}, Path: {file_path}")
    
    text = ""
    extraction_method_used = "Standard"
    
    try:
        if "pdf" in file_type.lower():
            logger.info("Processing PDF file...")
            
            # Try PyMuPDF for better PDF handling
            if PYMUPDF_AVAILABLE:
                logger.info("Using PyMuPDF for PDF processing")
                
                pdf_document = fitz.open(file_path)
                page_count = len(pdf_document)
                logger.info(f"PDF has {page_count} pages")
                
                # Try text extraction first
                for page_num in range(page_count):
                    page = pdf_document.load_page(page_num)
                    page_text = page.get_text()
                    text += page_text + "\n"
                
                logger.info(f"PyMuPDF text extraction: {len(text)} characters")
                extraction_method_used = "PyMuPDF Text Extraction"
                
                # If minimal text, try OCR
                if len(text.strip()) < 100:
                    logger.warning("Minimal text from PyMuPDF, using OCR on PDF pages")
                    text = ""  # Reset
                    extraction_method_used = "PyMuPDF + OCR"
                    
                    for page_num in range(page_count):
                        page = pdf_document.load_page(page_num)
                        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))  # Higher resolution
                        img_path = os.path.join(TEMP_DIR, f"temp_page_{page_num}.png")
                        pix.save(img_path)
                        
                        image = Image.open(img_path)
                        logger.info(f"Running OCR on page {page_num + 1}")
                        page_text = pytesseract.image_to_string(image, config='--psm 6')
                        text += page_text + "\n"
                        logger.info(f"OCR extracted {len(page_text)} characters from page {page_num + 1}")
                        
                        # Clean up temp file
                        os.remove(img_path)
                
                # If still minimal text, fallback to OpenAI Vision
                if len(text.strip()) < 200:
                    logger.warning("OCR results poor, falling back to OpenAI Vision API")
                    text = ""  # Reset
                    extraction_method_used = "OpenAI Vision API"
                    
                    client = setup_openai_client()
                    
                    for page_num in range(page_count):
                        page = pdf_document.load_page(page_num)
                        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
                        img_path = os.path.join(TEMP_DIR, f"temp_page_{page_num}.png")
                        pix.save(img_path)
                        
                        # Convert to base64 for OpenAI
                        with open(img_path, "rb") as img_file:
                            img_data = img_file.read()
                        img_base64 = base64.b64encode(img_data).decode('utf-8')
                        
                        logger.info(f"Using OpenAI Vision API for page {page_num + 1}")
                        page_text = await extract_text_with_openai_vision(img_base64, page_num + 1, client)
                        if page_text:
                            text += page_text + "\n"
                            logger.info(f"OpenAI Vision extracted {len(page_text)} characters from page {page_num + 1}")
                        
                        # Clean up temp file
                        os.remove(img_path)
                
                pdf_document.close()
                logger.info(f"Final PDF extraction completed: {len(text)} characters using {extraction_method_used}")
            
            elif PYPDF2_AVAILABLE:
                # Fallback to PyPDF2
                logger.warning("PyMuPDF not available, falling back to PyPDF2 + OCR")
                
                pdf_reader = PyPDF2.PdfReader(file_path)
                page_count = len(pdf_reader.pages)
                logger.info(f"PDF has {page_count} pages (PyPDF2 fallback)")
                
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += page_text + "\n"
                
                logger.info(f"PyPDF2 extraction completed: {len(text)} characters")
                extraction_method_used = "PyPDF2 (Basic)"
                
                # If minimal text with PyPDF2, warn about installing PyMuPDF
                if len(text.strip()) < 100:
                    logger.error("Minimal text extraction with PyPDF2. Better results require PyMuPDF.")
            else:
                logger.error("No PDF processing libraries available. Install PyMuPDF or PyPDF2.")
                raise HTTPException(
                    status_code=500, 
                    detail="PDF processing libraries not available. Install PyMuPDF or PyPDF2."
                )
                
        elif "docx" in file_type.lower() and DOCX_AVAILABLE:
            logger.info("Processing DOCX file...")
            doc = docx.Document(file_path)
            paragraph_count = len(doc.paragraphs)
            logger.info(f"DOCX has {paragraph_count} paragraphs")
            
            for i, paragraph in enumerate(doc.paragraphs):
                text += paragraph.text + "\n"
                
            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
            extraction_method_used = "DOCX Extraction"
                
        elif any(img_type in file_type.lower() for img_type in ["jpeg", "jpg", "png", "tiff", "image"]):
            logger.info(f"Processing image file with OCR: {file_type}")
            image = Image.open(file_path)
            logger.info(f"Image dimensions: {image.size}")
            
            text = pytesseract.image_to_string(image)
            logger.info(f"OCR extraction completed: {len(text)} characters extracted")
            extraction_method_used = "OCR (Tesseract)"
            
            # If minimal text from OCR, try OpenAI Vision
            if len(text.strip()) < 100:
                logger.warning("Minimal text from OCR, trying OpenAI Vision")
                client = setup_openai_client()
                
                with open(file_path, "rb") as img_file:
                    img_data = img_file.read()
                img_base64 = base64.b64encode(img_data).decode('utf-8')
                
                vision_text = await extract_text_with_openai_vision(img_base64, 0, client)
                if vision_text and len(vision_text.strip()) > len(text.strip()):
                    text = vision_text
                    extraction_method_used = "OpenAI Vision API"
                    logger.info(f"OpenAI Vision extracted {len(text)} characters")
            
        elif "text" in file_type.lower():
            logger.info("Processing text file...")
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                text = f.read()
            logger.info(f"Successfully read {len(text)} characters from text file")
            extraction_method_used = "Text File Read"
            
        else:
            error_msg = f"Unsupported file type: {file_type}"
            logger.error(error_msg)
            raise HTTPException(status_code=400, detail=error_msg)
        
        if len(text.strip()) == 0:
            logger.warning("No text content extracted from file")
            raise HTTPException(status_code=400, detail="No text content found in the uploaded file")
        
        logger.info(f"Text extraction completed successfully. Total characters: {len(text)}")
        return text, extraction_method_used
            
    except Exception as e:
        error_msg = f"Error extracting text from {file_name}: {str(e)}"
        logger.error(error_msg)
        raise HTTPException(status_code=500, detail=error_msg)

async def extract_text_with_openai_vision(img_base64: str, page_num: int, client=None) -> Optional[str]:
    """Extract text from image using OpenAI Vision API"""
    try:
        logger.info(f"Using OpenAI Vision API to extract text from page {page_num}")
        
        # Use the newer client interface if provided
        if client is not None:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": """Extract all the text from this document image. 
                                Maintain the original formatting and structure as much as possible.
                                Include all text, numbers, dates, and any other readable content.
                                Return only the extracted text without any commentary."""
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{img_base64}",
                                    "detail": "high"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=2000,
                temperature=0.1
            )
            
            extracted_text = response.choices[0].message.content
            
        else:
            # Fallback to older API style
            logger.info("Using legacy OpenAI API for vision")
            response = openai.ChatCompletion.create(
                model="gpt-4-vision-preview",  # Use vision model
                messages=[
                    {
                        "role": "user", 
                        "content": [
                            {
                                "type": "text",
                                "text": """Extract all the text from this document image. 
                                Maintain the original formatting and structure as much as possible.
                                Include all text, numbers, dates, and any other readable content.
                                Return only the extracted text without any commentary."""
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{img_base64}",
                                    "detail": "high"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=2000,
                temperature=0.1
            )
            
            extracted_text = response.choices[0].message.content
        
        logger.info(f"OpenAI Vision extracted {len(extracted_text)} characters from page {page_num}")
        return extracted_text
        
    except Exception as e:
        error_msg = f"Error using OpenAI Vision API for page {page_num}: {str(e)}"
        logger.error(error_msg)
        return None

async def extract_contract_info(document_text: str, document_type: str = "Rental") -> Optional[Dict[str, Any]]:
    """Use GPT-4o-mini to extract comprehensive contract information with document-specific prompts"""
    logger.info(f"Starting contract information extraction with GPT-4o-mini for {document_type}")
    logger.info(f"Input text length: {len(document_text)} characters")
    
    try:
        # Get document-specific schema and prompts
        schema = get_extraction_schema(document_type)
        
        # Create document-specific system prompts
        if document_type == "NDA":
            system_prompt = f"""You are an expert legal contract analyzer specializing in Non-Disclosure Agreements (NDAs).
            
            Extract comprehensive information from this NDA document and return a JSON object with the following structure:
            {schema}
            
            Pay special attention to:
            - Disclosing and receiving parties
            - Confidentiality scope and definitions
            - Duration of confidentiality obligations
            - Restrictions and exceptions
            - Survival periods after termination
            
            If any information is not found, use "Not specified" as the value.
            Ensure dates are in YYYY-MM-DD format for calculation purposes.
            Make sure document_type is set to "Non-Disclosure Agreement"."""
            
            user_prompt = f"Extract information from this Non-Disclosure Agreement:\n\n{document_text}"
            
        elif document_type == "MSA":
            system_prompt = f"""You are an expert legal contract analyzer specializing in Master Service Agreements (MSAs).
            
            Extract comprehensive information from this MSA document and return a JSON object with the following structure:
            {schema}
            
            Pay special attention to:
            - Service provider and client details
            - Description of services and deliverables
            - Payment terms and pricing models
            - Contract duration and renewal terms
            - Termination conditions and effects
            - Key obligations for each party
            
            If any information is not found, use "Not specified" as the value.
            Ensure dates are in YYYY-MM-DD format for calculation purposes.
            Make sure document_type is set to "Master Service Agreement"."""
            
            user_prompt = f"Extract information from this Master Service Agreement:\n\n{document_text}"
            
        elif document_type == "Insurance":
            system_prompt = f"""You are an expert insurance policy analyzer specializing in all types of insurance policies.
            
            Extract comprehensive information from this insurance policy document and return a JSON object with the following structure:
            {schema}
            
            Pay special attention to:
            - Insurance company and policyholder details
            - Policy type (Health, Auto, Life, Property, etc.) and coverage amounts
            - Premium costs, deductibles, copayments, and coinsurance
            - Policy effective dates and renewal information
            - What is covered and what is excluded
            - Claim procedures and requirements
            
            If any information is not found, use "Not specified" as the value.
            Ensure dates are in YYYY-MM-DD format for calculation purposes.
            Make sure document_type is set to "Insurance Policy"."""
            
            user_prompt = f"Extract information from this Insurance Policy:\n\n{document_text}"
            
        elif document_type == "MOU":
            system_prompt = f"""You are an expert contract analyzer specializing in Memorandums of Understanding (MOUs).
            
            Extract comprehensive information from this MOU document and return a JSON object with the following structure:
            {schema}
            
            Pay special attention to:
            - All parties involved (bilateral or multilateral)
            - Purpose, objectives, and scope of the MOU
            - Specific responsibilities and obligations of each party
            - Duration and review periods
            - Governance and management structure
            - Termination conditions
            
            If any information is not found, use "Not specified" as the value.
            Ensure dates are in YYYY-MM-DD format for calculation purposes.
            Make sure document_type is set to "Memorandum of Understanding"."""
            
            user_prompt = f"Extract information from this Memorandum of Understanding:\n\n{document_text}"
            
        else:  # Default to Rental
            system_prompt = f"""You are an expert legal contract analyzer specializing in rental/lease agreements.
            
            Extract comprehensive information from this rental document and return a JSON object with the following structure:
            {schema}
            
            Pay special attention to:
            - Landlord and tenant details
            - Property address and description
            - Rent amounts and financial terms
            - Lease start and end dates
            - Key terms and special clauses
            
            If any information is not found, use "Not specified" as the value.
            Ensure dates are in YYYY-MM-DD format for calculation purposes.
            Make sure document_type is set to "Rental Agreement"."""
            
            user_prompt = f"Extract information from this rental/lease document:\n\n{document_text}"
        
        logger.info(f"Using {document_type}-specific extraction prompt")
        logger.info("Sending request to OpenAI API...")
        
        # Try to use the newer OpenAI client
        client = setup_openai_client()
        
        if client is not None:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": system_prompt
                    },
                    {
                        "role": "user",
                        "content": user_prompt
                    }
                ],
                temperature=0.1,
                max_tokens=2000
            )
            
            content = response.choices[0].message.content
            
        else:
            # Fallback to legacy API
            response = openai.ChatCompletion.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": system_prompt
                    },
                    {
                        "role": "user",
                        "content": user_prompt
                    }
                ],
                temperature=0.1,
                max_tokens=2000
            )
            
            content = response.choices[0].message.content
        
        logger.info("Received response from OpenAI API")
        logger.info(f"Response content length: {len(content)} characters")
        
        # Try to parse JSON
        try:
            # First, try to extract JSON from markdown code blocks if present
            json_content = content
            
            # Check if content is wrapped in markdown code blocks
            if '```json' in content:
                logger.info("Detected JSON wrapped in markdown code blocks, extracting...")
                # Extract content between ```json and ```
                json_match = re.search(r'```json\s*\n(.*?)\n```', content, re.DOTALL)
                if json_match:
                    json_content = json_match.group(1)
                    logger.info("Successfully extracted JSON from markdown wrapper")
                else:
                    # Try alternative pattern
                    json_match = re.search(r'```json(.*?)```', content, re.DOTALL)
                    if json_match:
                        json_content = json_match.group(1).strip()
                        logger.info("Successfully extracted JSON using alternative pattern")
            
            # Now try to parse the cleaned JSON
            parsed_data = json.loads(json_content)
            logger.info("Successfully parsed JSON response from OpenAI")
            
            return parsed_data
            
        except json.JSONDecodeError as e:
            logger.warning(f"Failed to parse JSON response: {str(e)}")
            logger.warning(f"Attempted to parse: {json_content[:200]}...")
            
            # Try manual extraction as fallback
            logger.info("Attempting manual information extraction as fallback")
            fallback_data = {
                "contract_summary": "Contract analysis completed but JSON parsing failed",
                "raw_analysis": content,
                "extraction_error": f"JSON parsing failed: {str(e)}",
                "document_type": document_type
            }
            return fallback_data
            
    except Exception as e:
        error_msg = f"Error extracting contract info: {str(e)}"
        logger.error(error_msg)
        
        # Return basic error info
        return {
            "document_type": document_type,
            "extraction_error": error_msg,
            "contract_summary": "Error occurred during contract analysis"
        }

def calculate_days_until_expiry(end_date_str: str) -> tuple:
    """Calculate days until contract expiry"""
    logger.info(f"Calculating days until expiry for date: {end_date_str}")
    
    try:
        if end_date_str and end_date_str != "Not specified":
            end_date = parser.parse(end_date_str)
            today = datetime.now()
            days_left = (end_date - today).days
            
            logger.info(f"Contract end date: {end_date.strftime('%Y-%m-%d')}")
            logger.info(f"Today's date: {today.strftime('%Y-%m-%d')}")
            logger.info(f"Days until expiry: {days_left}")
            
            return days_left, end_date
        else:
            logger.warning("End date not specified or invalid")
            return None, None
    except Exception as e:
        logger.error(f"Error calculating days until expiry: {str(e)}")
        return None, None

def create_brief_summary(data: Dict[str, Any], document_type: str) -> str:
    """Create a brief 3-4 line summary of the contract"""
    
    if "Non-Disclosure" in document_type or "NDA" in document_type:
        parties = data.get("parties", {})
        disclosing = parties.get("disclosing_party", "Party 1")
        receiving = parties.get("receiving_party", "Party 2")
        dates = data.get("dates", {})
        duration = dates.get("duration", "specified period")
        
        return f"""This Non-Disclosure Agreement is between {disclosing} (disclosing party) and {receiving} (receiving party). The agreement protects confidential information shared for business discussions and evaluation purposes. Confidentiality obligations remain in effect for {duration}. Both parties must keep shared information confidential and use it only for the intended business purpose."""
    
    elif "Master Service" in document_type or "MSA" in document_type:
        parties = data.get("parties", {})
        provider = parties.get("service_provider", "Service Provider")
        client = parties.get("client", "Client")
        service_details = data.get("service_details", {})
        services = service_details.get("description", "professional services")
        dates = data.get("dates", {})
        term = dates.get("initial_term", "specified period")
        
        return f"""This Master Service Agreement is between {provider} (service provider) and {client} (client). The provider will deliver {services} under the terms and conditions outlined in the agreement. The contract has an initial term of {term} and covers service delivery, payment terms, and responsibilities of both parties."""
    
    elif "Insurance" in document_type:
        parties = data.get("parties", {})
        insurer = parties.get("insurer", "Insurance Company")
        policyholder = parties.get("policyholder", "Policyholder")
        policy_details = data.get("policy_details", {})
        insurance_type = policy_details.get("insurance_type", "insurance coverage")
        coverage_amount = policy_details.get("coverage_amount", "specified coverage")
        financial = data.get("financial_terms", {})
        premium = financial.get("premium_amount", "agreed premium")
        
        return f"""This insurance policy is between {insurer} (insurer) and {policyholder} (policyholder). The policy provides {insurance_type} with {coverage_amount} coverage. The policyholder pays {premium} in premiums and the policy outlines covered benefits, exclusions, and claim procedures."""
    
    elif "Memorandum" in document_type or "MOU" in document_type:
        parties = data.get("parties", {})
        party_1 = parties.get("party_1", "First Party")
        party_2 = parties.get("party_2", "Second Party")
        purpose = data.get("purpose", {})
        objective = purpose.get("objective", "collaborative objectives")
        dates = data.get("dates", {})
        duration = dates.get("duration", "specified period")
        
        return f"""This Memorandum of Understanding is between {party_1} and {party_2} for {objective}. The MOU establishes a framework for cooperation and collaboration between the parties. The agreement lasts for {duration} and outlines the responsibilities, governance structure, and terms of the partnership."""
    
    else:  # Rental
        parties = data.get("parties", {})
        landlord = parties.get("landlord", "Landlord")
        tenant = parties.get("tenant", "Tenant")
        property_details = data.get("property_details", {})
        address = property_details.get("address", "specified property")
        financial = data.get("financial_terms", {})
        rent = financial.get("monthly_rent", "agreed amount")
        dates = data.get("dates", {})
        term = dates.get("lease_term", "specified period")
        
        return f"""This rental agreement is between {landlord} (landlord) and {tenant} (tenant) for the property located at {address}. The lease term is {term} with monthly rent of {rent}. The agreement outlines the rights, responsibilities, and obligations of both landlord and tenant during the rental period."""

async def answer_question_locally(question: str, contract_data: Dict[str, Any]) -> Optional[str]:
    """Answer simple questions directly from structured data (instant responses for all document types)"""
    start_time = time.perf_counter()
    
    question_lower = question.lower()
    document_type = contract_data.get("document_type", "Unknown")
    logger.info(f"Attempting local answer for {document_type}: {question}")
    
    try:
        # Universal WHO questions (works for all document types)
        if any(word in question_lower for word in ["who", "party", "parties"]):
            parties = contract_data.get("parties", {})
            
            # Rental-specific
            if "tenant" in question_lower or "lessee" in question_lower:
                tenant = parties.get("tenant", "Not specified")
                if tenant != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (tenant): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The person renting the property is: {tenant}"
            
            if "landlord" in question_lower or "lessor" in question_lower:
                landlord = parties.get("landlord", "Not specified")
                if landlord != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (landlord): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The property owner is: {landlord}"
            
            # NDA-specific
            if "disclosing" in question_lower:
                disclosing = parties.get("disclosing_party", "Not specified")
                if disclosing != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (disclosing party): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The company sharing confidential information is: {disclosing}"
            
            if "receiving" in question_lower:
                receiving = parties.get("receiving_party", "Not specified")
                if receiving != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (receiving party): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The company receiving confidential information is: {receiving}"
            
            # MSA-specific
            if "provider" in question_lower or "service" in question_lower and "provider" in question_lower:
                provider = parties.get("service_provider", "Not specified")
                if provider != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (service provider): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The company providing the services is: {provider}"
            
            if "client" in question_lower:
                client = parties.get("client", "Not specified")
                if client != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (client): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The company buying the services is: {client}"
            
            # Insurance-specific
            if "insurer" in question_lower or "insurance company" in question_lower:
                insurer = parties.get("insurer", "Not specified")
                if insurer != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (insurer): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The insurance company is: {insurer}"
            
            if "policyholder" in question_lower or "insured" in question_lower:
                policyholder = parties.get("policyholder", "Not specified")
                if policyholder != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (policyholder): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The person being insured is: {policyholder}"
            
            # MOU-specific
            if "first party" in question_lower or "party 1" in question_lower:
                party_1 = parties.get("party_1", "Not specified")
                if party_1 != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (party 1): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The first party is: {party_1}"
            
            if "second party" in question_lower or "party 2" in question_lower:
                party_2 = parties.get("party_2", "Not specified")
                if party_2 != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (party 2): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The second party is: {party_2}"
        
        # Universal WHEN questions (works for all document types)
        if any(word in question_lower for word in ["when", "expire", "end", "start", "date"]):
            dates = contract_data.get("dates", {})
            
            if "expire" in question_lower or "end" in question_lower:
                # For rentals, use the direct end_date
                if "Rental" in document_type:
                    end_date = dates.get("end_date", "Not specified")
                    if end_date != "Not specified":
                        try:
                            end_dt = parser.parse(end_date)
                            today = datetime.now()
                            days_left = (end_dt - today).days
                            
                            if days_left > 0:
                                return f"The lease expires on {end_date} ({days_left} days remaining)."
                            else:
                                return f"The lease expired on {end_date} ({abs(days_left)} days ago)."
                        except:
                            return f"The lease expires/ends on {end_date}."
                
                # Generic fallback for any document type
                end_date = dates.get("end_date") or dates.get("policy_end") or "Not specified"
                if end_date != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (generic expiry): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The agreement runs out on {end_date}."
            
            if "start" in question_lower or "effective" in question_lower:
                start_date = (dates.get("start_date") or 
                             dates.get("effective_date") or 
                             dates.get("execution_date") or 
                             dates.get("policy_start") or 
                             "Not specified")
                if start_date != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (start): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The agreement started on {start_date}."
        
        # Document-specific financial questions
        if any(word in question_lower for word in ["how much", "cost", "price", "amount", "money", "pay"]):
            
            if "Rental" in document_type:
                financial = contract_data.get("financial_terms", {})
                if "rent" in question_lower:
                    rent = financial.get("monthly_rent", "Not specified")
                    if rent != "Not specified":
                        end_time = time.perf_counter()
                        logger.info(f"⚡ Local answer (rent): {(end_time - start_time) * 1000:.2f}ms")
                        return f"The monthly rent is {rent}."
                
                if "deposit" in question_lower:
                    deposit = financial.get("security_deposit", "Not specified")
                    if deposit != "Not specified":
                        end_time = time.perf_counter()
                        logger.info(f"⚡ Local answer (deposit): {(end_time - start_time) * 1000:.2f}ms")
                        return f"The security deposit (upfront payment) is {deposit}."
            
            elif "MSA" in document_type:
                financial = contract_data.get("financial_terms", {})
                if "payment" in question_lower:
                    payment_terms = financial.get("payment_terms", "Not specified")
                    if payment_terms != "Not specified":
                        end_time = time.perf_counter()
                        logger.info(f"⚡ Local answer (payment terms): {(end_time - start_time) * 1000:.2f}ms")
                        return f"Here's how payments work: {payment_terms}."
            
            elif "Insurance" in document_type:
                financial = contract_data.get("financial_terms", {})
                if "premium" in question_lower:
                    premium = financial.get("premium_amount", "Not specified")
                    if premium != "Not specified":
                        end_time = time.perf_counter()
                        logger.info(f"⚡ Local answer (premium): {(end_time - start_time) * 1000:.2f}ms")
                        return f"The insurance premium is {premium}."
                
                if "deductible" in question_lower:
                    policy_details = contract_data.get("policy_details", {})
                    deductible = policy_details.get("deductible", "Not specified")
                    if deductible != "Not specified":
                        end_time = time.perf_counter()
                        logger.info(f"⚡ Local answer (deductible): {(end_time - start_time) * 1000:.2f}ms")
                        return f"The deductible amount is {deductible}."
        
        # Location questions (mainly for rental)
        if any(word in question_lower for word in ["where", "address", "location", "property"]):
            if "Rental" in document_type:
                property_details = contract_data.get("property_details", {})
                address = property_details.get("address", "Not specified")
                if address != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (address): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The property is located at: {address}"
        
        # Duration/term questions (universal)
        if any(word in question_lower for word in ["duration", "term", "how long", "period"]):
            dates = contract_data.get("dates", {})
            
            # Try different term fields based on document type
            term = (dates.get("lease_term") or 
                   dates.get("initial_term") or 
                   dates.get("duration") or 
                   "Not specified")
            
            if term != "Not specified":
                end_time = time.perf_counter()
                logger.info(f"⚡ Local answer (term): {(end_time - start_time) * 1000:.2f}ms")
                return f"The agreement lasts for {term}."
        
        # If no local answer found
        end_time = time.perf_counter()
        logger.info(f"📊 No local answer found: {(end_time - start_time) * 1000:.2f}ms")
        return None
        
    except Exception as e:
        logger.error(f"Error in local answering: {str(e)}")
        return None

async def chat_with_document(question: str, contract_data: Dict[str, Any], document_text: str) -> str:
    """Hybrid chat: Try local answer first, fallback to full document text with OpenAI for complex questions"""
    overall_start = time.perf_counter()
    logger.info(f"=== HYBRID CHAT ANALYSIS START ===")
    logger.info(f"Question: {question}")
    
    # Step 1: Try local answer first (instant)
    local_start = time.perf_counter()
    local_answer = await answer_question_locally(question, contract_data)
    local_end = time.perf_counter()
    
    if local_answer:
        total_time = local_end - overall_start
        logger.info(f"📊 ⚡ INSTANT LOCAL ANSWER: {total_time * 1000:.2f}ms")
        logger.info(f"📊 Answer: {local_answer}")
        logger.info(f"=== HYBRID CHAT ANALYSIS END ===")
        return local_answer
    
    # Step 2: Fallback to OpenAI using FULL DOCUMENT TEXT for complex questions
    logger.info("📊 Local answer not found, using OpenAI with full document text")
    
    try:
        # Context preparation using FULL document text + extracted data
        context_start = time.perf_counter()
        
        # Use full document text for comprehensive answers
        full_context = f"""
        FULL DOCUMENT TEXT:
        {document_text}
        
        EXTRACTED STRUCTURED DATA:
        {json.dumps(contract_data, indent=2)}
        """
        
        context_end = time.perf_counter()
        logger.info(f"📊 Full context preparation: {(context_end - context_start) * 1000:.2f}ms")
        logger.info(f"📊 Full context size: {len(full_context)} characters")
        
        # Client setup
        client = setup_openai_client()
        document_type = contract_data.get("document_type", "Contract")
        
        # Message preparation
        messages = [
            {
                "role": "system",
                "content": f"""You are a helpful assistant that explains {document_type} documents in simple, easy-to-understand language.
                
                You have access to:
                1. The FULL DOCUMENT TEXT - Use this for detailed questions about specific clauses, terms, conditions, obligations, or any content in the document
                2. EXTRACTED STRUCTURED DATA - Use this for quick reference to key information
                
                Instructions:
                - Answer questions in simple, everyday language that anyone can understand
                - Avoid legal jargon - use plain English instead
                - Explain things like you're talking to a friend or family member
                - For specific clauses or detailed terms, quote relevant sections from the document but explain what they mean in simple terms
                - Be accurate but make it easy to understand
                - If something is complex, break it down into simple steps or bullet points
                - Use analogies or examples when helpful
                - If information is not in the document, clearly state that"""
            },
            {
                "role": "user", 
                "content": f"Document Context: {full_context}\n\nQuestion: {question}"
            }
        ]
        
        # API call
        api_start = time.perf_counter()
        logger.info("📊 Sending request to OpenAI API for comprehensive document analysis...")
        
        if client is not None:
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=messages,
                temperature=0.1,
                max_tokens=500
            )
            answer = response.choices[0].message.content
        else:
            # Fallback to legacy API
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=messages,
                temperature=0.1,
                max_tokens=500,
                request_timeout=10
            )
            answer = response.choices[0].message.content
        
        api_end = time.perf_counter()
        api_time = api_end - api_start
        logger.info(f"📊 ⚡ OpenAI API call: {api_time * 1000:.2f}ms")
        
        # Overall timing
        overall_end = time.perf_counter()
        total_time = overall_end - overall_start
        local_time = local_end - local_start
        
        logger.info(f"📊 ⚡ TOTAL HYBRID CHAT TIME: {total_time * 1000:.2f}ms")
        logger.info(f"📊 Local attempt: {local_time * 1000:.2f}ms")
        logger.info(f"📊 API time: {api_time * 1000:.2f}ms")
        logger.info(f"=== HYBRID CHAT ANALYSIS END ===")
        
        return answer
        
    except Exception as e:
        error_msg = f"Error in OpenAI fallback: {str(e)}"
        logger.error(f"📊 ❌ HYBRID CHAT ERROR: {error_msg}")
        return f"An error occurred while processing your question: {str(e)}"

# Background Tasks
async def process_document(document_id: str, file_path: str, filename: str, file_type: str, document_type: str):
    """Background task to process document"""
    try:
        logger.info(f"Starting background processing for document: {document_id}")
        
        # Extract text from document
        document_text, extraction_method = await extract_text_from_file(file_path, file_type, filename)
        
        if not document_text:
            logger.error(f"Text extraction failed for document: {document_id}")
            ANALYSIS_RESULTS[document_id] = {
                "document_id": document_id,
                "document_type": document_type,
                "status": "error",
                "error": "Text extraction failed",
                "analysis_complete": False
            }
            return
        
        # Store document text
        DOCUMENT_STORE[document_id] = {
            "text": document_text,
            "filename": filename,
            "document_type": document_type,
            "extraction_method": extraction_method,
            "char_count": len(document_text),
            "word_count": len(document_text.split())
        }
        
        logger.info(f"Text extraction completed for {document_id}. Starting analysis...")
        
        # Analyze document with OpenAI
        contract_data = await extract_contract_info(document_text, document_type)
        
        if contract_data:
            # Add document metadata
            contract_data["document_id"] = document_id
            contract_data["extraction_method"] = extraction_method
            
            # Create brief summary
            if "contract_summary" not in contract_data or not contract_data["contract_summary"]:
                document_type = contract_data.get("document_type", document_type)
                contract_data["contract_summary"] = create_brief_summary(contract_data, document_type)
            
            # Store analysis results
            ANALYSIS_RESULTS[document_id] = {
                "document_id": document_id,
                "document_type": contract_data.get("document_type", document_type),
                "data": contract_data,
                "status": "completed",
                "extraction_method": extraction_method,
                "analysis_complete": True
            }
            
            logger.info(f"Analysis completed for document: {document_id}")
        else:
            logger.error(f"Analysis failed for document: {document_id}")
            ANALYSIS_RESULTS[document_id] = {
                "document_id": document_id,
                "document_type": document_type,
                "status": "error",
                "error": "Analysis failed",
                "analysis_complete": False
            }
        
        # Clean up temporary file
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"Removed temporary file: {file_path}")
        except Exception as e:
            logger.warning(f"Failed to remove temporary file {file_path}: {str(e)}")
            
    except Exception as e:
        logger.error(f"Error processing document {document_id}: {str(e)}")
        ANALYSIS_RESULTS[document_id] = {
            "document_id": document_id,
            "document_type": document_type,
            "status": "error",
            "error": str(e),
            "analysis_complete": False
        }

# API Endpoints
@app.get("/", tags=["General"])
async def root():
    """Root endpoint"""
    return {
        "name": "Multi-Document Contract Analyzer API",
        "version": "1.0.0",
        "description": "API for analyzing various contract types and providing Q&A capability"
    }

@app.post("/upload", response_model=DocumentUploadResponse, tags=["Document Management"])
async def upload_document(
    background_tasks: BackgroundTasks,
    document_type: str = Form(...),
    file: UploadFile = File(...)
):
    """Upload a contract document for processing"""
    try:
        # Check if OpenAI API key is configured
        if not OPENAI_API_KEY:
            raise HTTPException(status_code=500, detail="OpenAI API key not configured")
        
        # Check file type
        file_type = file.content_type
        allowed_types = ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         "text/plain", "image/jpeg", "image/jpg", "image/png", "image/tiff"]
        
        if file_type not in allowed_types:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_type}")
        
        # Generate unique document ID
        document_id = str(uuid.uuid4())
        filename = file.filename
        
        # Check document type
        valid_document_types = ["Rental", "NDA", "MSA", "Insurance", "MOU"]
        if document_type not in valid_document_types:
            raise HTTPException(status_code=400, detail=f"Invalid document type: {document_type}. Supported types: {', '.join(valid_document_types)}")
        
        # Create directory for temporary files if it doesn't exist
        os.makedirs(TEMP_DIR, exist_ok=True)
        
        # Save the file to a temporary location
        file_path = os.path.join(TEMP_DIR, f"{document_id}_{filename}")
        with open(file_path, "wb") as temp_file:
            shutil.copyfileobj(file.file, temp_file)
        
        # Add to processing queue
        ANALYSIS_RESULTS[document_id] = {
            "document_id": document_id,
            "document_type": document_type,
            "status": "processing",
            "analysis_complete": False
        }
        
        # Schedule background processing
        background_tasks.add_task(process_document, document_id, file_path, filename, file_type, document_type)
        
        logger.info(f"Document queued for processing: {document_id}")
        
        return {
            "document_id": document_id,
            "filename": filename,
            "document_type": document_type,
            "status": "processing",
            "message": "Document uploaded and queued for processing"
        }
        
    except Exception as e:
        logger.error(f"Error uploading document: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/documents/{document_id}", response_model=DocumentInfoResponse, tags=["Document Management"])
async def get_document_info(document_id: str):
    """Get information about a processed document"""
    if document_id not in ANALYSIS_RESULTS:
        raise HTTPException(status_code=404, detail=f"Document not found: {document_id}")
    
    result = ANALYSIS_RESULTS[document_id]
    document_info = DOCUMENT_STORE.get(document_id, {})
    
    return {
        "document_id": document_id,
        "document_type": result.get("document_type", "Unknown"),
        "filename": document_info.get("filename", "Unknown"),
        "extraction_method": document_info.get("extraction_method", "Standard"),
        "char_count": document_info.get("char_count", 0),
        "word_count": document_info.get("word_count", 0),
        "status": result.get("status", "unknown"),
        "analysis_complete": result.get("analysis_complete", False)
    }

@app.get("/analysis/{document_id}", tags=["Document Analysis"])
async def get_document_analysis(document_id: str):
    """Get analysis results for a document"""
    if document_id not in ANALYSIS_RESULTS:
        raise HTTPException(status_code=404, detail=f"Document not found: {document_id}")
    
    result = ANALYSIS_RESULTS[document_id]
    
    if result["status"] == "processing":
        return {
            "document_id": document_id,
            "status": "processing",
            "message": "Document is still being processed"
        }
    
    if result["status"] == "error":
        return {
            "document_id": document_id,
            "status": "error",
            "error": result.get("error", "Unknown error")
        }
    
    # Return the full analysis data
    return result.get("data", {})

@app.post("/question", response_model=QuestionResponse, tags=["Document Q&A"])
async def ask_question(question_req: QuestionRequest):
    """Ask a question about a document"""
    document_id = question_req.document_id
    question = question_req.question
    session_id = question_req.session_id or str(uuid.uuid4())
    
    logger.info(f"Question received for document {document_id}: {question}")
    
    # Validate document exists and has been processed
    if document_id not in ANALYSIS_RESULTS:
        raise HTTPException(status_code=404, detail=f"Document not found: {document_id}")
    
    result = ANALYSIS_RESULTS[document_id]
    if result["status"] != "completed" or not result.get("analysis_complete", False):
        raise HTTPException(status_code=400, detail="Document analysis is not complete")
    
    # Get document text and analysis data
    if document_id not in DOCUMENT_STORE:
        raise HTTPException(status_code=404, detail="Document text not found")
    
    document_text = DOCUMENT_STORE[document_id]["text"]
    contract_data = result["data"]
    
    # Process start time for performance tracking
    start_time = time.perf_counter()
    
    # Get answer to question
    answer = await chat_with_document(question, contract_data, document_text)
    
    # Calculate processing time
    end_time = time.perf_counter()
    processing_time_ms = (end_time - start_time) * 1000
    
    # Store question/answer in session history
    if session_id not in SESSIONS:
        SESSIONS[session_id] = []
    
    SESSIONS[session_id].append({
        "question": question,
        "answer": answer,
        "timestamp": datetime.now().isoformat()
    })
    
    logger.info(f"Answered question for document {document_id} in {processing_time_ms:.2f}ms")
    
    return {
        "document_id": document_id,
        "question": question,
        "answer": answer,
        "session_id": session_id,
        "processing_time_ms": processing_time_ms
    }

@app.get("/documents", tags=["Document Management"])
async def list_documents():
    """List all documents in the system"""
    documents = []
    
    for doc_id, result in ANALYSIS_RESULTS.items():
        document_info = DOCUMENT_STORE.get(doc_id, {})
        
        documents.append({
            "document_id": doc_id,
            "document_type": result.get("document_type", "Unknown"),
            "filename": document_info.get("filename", "Unknown"),
            "status": result.get("status", "unknown"),
            "analysis_complete": result.get("analysis_complete", False),
            "char_count": document_info.get("char_count", 0),
            "timestamp": document_info.get("timestamp", datetime.now().isoformat())
        })
    
    return {
        "count": len(documents),
        "documents": documents
    }

@app.get("/sessions/{session_id}", tags=["Document Q&A"])
async def get_session_history(session_id: str):
    """Get conversation history for a session"""
    if session_id not in SESSIONS:
        raise HTTPException(status_code=404, detail=f"Session not found: {session_id}")
    
    return {
        "session_id": session_id,
        "conversation_history": SESSIONS[session_id]
    }

@app.on_event("shutdown")
async def shutdown_event():
    """Clean up resources on shutdown"""
    # Remove temporary directory
    try:
        shutil.rmtree(TEMP_DIR, ignore_errors=True)
        logger.info(f"Removed temporary directory: {TEMP_DIR}")
    except Exception as e:
        logger.error(f"Error removing temporary directory: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    # Start the API server
    uvicorn.run(app, host="0.0.0.0", port=8000)